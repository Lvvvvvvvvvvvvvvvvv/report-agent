# -*- coding: utf-8 -*-
"""报告智能体核心引擎 v3 — 大模型生成模式（代码算数据，LLM写内容，模板降级兜底）"""
import os
from pathlib import Path
from datetime import datetime
import numpy as np, pandas as pd
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE=Path(__file__).parent; DATA_DIR=BASE/"data"
# 云端部署时 app 目录可能只读，自动回退到 /tmp
_proj_out = BASE/"output"
try:
    _proj_out.mkdir(exist_ok=True)
    (_proj_out/".w").write_text("ok"); (_proj_out/".w").unlink()
    OUT_DIR = _proj_out
except Exception:
    OUT_DIR = Path("/tmp/report_agent_out"); OUT_DIR.mkdir(exist_ok=True)
IMG_DIR = OUT_DIR/"_charts"; IMG_DIR.mkdir(exist_ok=True)
plt.rcParams["font.sans-serif"]=["Arial Unicode MS","PingFang SC","Heiti TC","STHeiti","SimHei"]
plt.rcParams["axes.unicode_minus"]=False

META={"unit":"三监区","parent":"监狱教育改造科","author":"","reviewer":"","approver":"","period":"2026年4月","cycle":"月度"}
F_TITLE,F_HEAD,F_BODY="黑体","黑体","仿宋"

# ── LLM 配置 ────────────────────────────────────────────
def _load_key():
    k=os.getenv("MIMO_API_KEY","").strip()
    if not k:
        f=BASE/"mimo_key.txt"
        if f.exists(): k=f.read_text(encoding="utf-8").strip()
    return "" if (not k or k.startswith("在这里")) else k

MIMO_KEY=_load_key()
MIMO_BASE_URL=os.getenv("MIMO_BASE_URL","https://token-plan-cn.xiaomimimo.com/anthropic")
MIMO_MODEL=os.getenv("MIMO_MODEL","mimo-v2.5")
LLM_STYLE="规范详实"
USE_LLM=bool(MIMO_KEY); _client=None
if USE_LLM:
    try:
        from anthropic import Anthropic
        _client=Anthropic(api_key=MIMO_KEY,base_url=MIMO_BASE_URL)
    except Exception as e:
        print("[警告] 大模型初始化失败：",e); USE_LLM=False

# 生成模式系统提示：LLM 接收结构化数据，从头生成公文段落
_SYS_GEN=(
    "你是公安监狱系统的资深数据分析师兼公文写手，专门撰写面向上级的行政公文报告。\n"
    "我会给你某章节的结构化统计数据，请根据这些数据直接生成该章节的公文正文段落。\n"
    "严格要求：\n"
    "①所有数字必须严格来自所给数据，不得虚构或引用不存在的数字；\n"
    "②语言风格规范庄重，符合监狱系统行政公文规范，使用第三人称；\n"
    "③直接输出正文内容，不加标题、不加编号前缀、不加任何解释说明；\n"
    "④长度100-200字，重点突出，逻辑清晰，措辞精准，有分析有结论有建议导向。"
)
_STYLE_MAP={
    "简洁精炼":"每段严格控制在100字以内，去套话，只保留核心数据和结论。",
    "严谨正式":"语气严谨正式，多用专业术语，符合行政公文最高规范。"
}
_SEC_LABELS={
    "overview":"报告概述","train_overall":"训练任务总体完成情况","train_cat":"分训练类别分析",
    "train_squad":"分中队训练对比","train_trend":"训练环比趋势","train_focus":"重点关注人员情况",
    "abn_overall":"异常行为总体态势","abn_risk":"异常行为风险等级分布",
    "abn_squad":"异常行为分中队分布","abn_high":"高风险个案详情",
    "abn_disposal":"异常行为处置与化解情况","conclusion":"综合研判与结论","suggestion":"下一步工作建议"
}

def _sys_gen(): return _SYS_GEN+("\n"+_STYLE_MAP[LLM_STYLE] if LLM_STYLE in _STYLE_MAP else "")

def _gen(key, data_str, fallback_fn, s):
    """核心生成函数：以结构化统计数据为输入，让LLM生成章节公文段落；失败自动降级模板。"""
    if not USE_LLM: return fallback_fn(s)
    msg=(f"章节名称：{_SEC_LABELS.get(key, key)}\n\n"
         f"统计数据如下：\n{data_str}\n\n"
         f"请根据以上数据，生成该章节的规范公文段落（100-200字）。")
    try:
        r=_client.messages.create(model=MIMO_MODEL,max_tokens=600,
            system=_sys_gen(),messages=[{"role":"user","content":msg}])
        result="".join(b.text for b in r.content if b.type=="text").strip()
        return result if result else fallback_fn(s)
    except Exception as e:
        print(f"  [警告] {key} 生成失败，降级模板：",e); return fallback_fn(s)


# ── ① 数据 ──────────────────────────────────────────────
def load_data():
    return pd.read_excel(DATA_DIR/"训练任务数据.xlsx"), pd.read_excel(DATA_DIR/"异常行为数据.xlsx")

# ── ② 分析 ──────────────────────────────────────────────
def analyze(train, abn, pass_score=60, pass_att=80, focus_threshold=70, squads=None):
    if squads:
        train=train[train["中队"].isin(squads)].copy(); abn=abn[abn["中队"].isin(squads)].copy()
    else: train=train.copy()
    train["是否达标"]=train.apply(lambda r:"是" if(r["考核成绩"]>=pass_score and r["出勤率(%)"]>=pass_att)else"否",axis=1)
    SQ=["一中队","二中队","三中队"]; aq=[q for q in SQ if q in train["中队"].unique()]
    s={}
    s["inmate_count"]=int(train["编号"].nunique()); s["train_records"]=int(len(train))
    s["avg_att"]=round(train["出勤率(%)"].mean(),1); s["pass_rate"]=round((train["是否达标"]=="是").mean()*100,1)
    s["avg_score"]=round(train["考核成绩"].mean(),1); s["pass_score"]=pass_score; s["pass_att"]=pass_att
    s["focus_threshold"]=focus_threshold
    s["by_cat"]=train.groupby("训练类别").agg(平均出勤率=("出勤率(%)","mean"),达标率=("是否达标",lambda x:(x=="是").mean()*100),平均成绩=("考核成绩","mean")).round(1)
    s["by_squad"]=train.groupby("中队").agg(平均出勤率=("出勤率(%)","mean"),达标率=("是否达标",lambda x:(x=="是").mean()*100),平均成绩=("考核成绩","mean")).round(1).reindex(aq)
    s["trend"]=train["较上月"].value_counts().to_dict()
    per=train.groupby(["编号","姓名","中队"]).agg(平均成绩=("考核成绩","mean"),平均出勤率=("出勤率(%)","mean"),不达标项=("是否达标",lambda x:int((x=="否").sum())),退步项=("较上月",lambda x:int((x=="退步").sum()))).round(1).reset_index()
    s["focus"]=per[(per["不达标项"]>=2)|(per["平均成绩"]<focus_threshold)].sort_values("平均成绩").head(6)
    s["abn_total"]=int(len(abn)); s["abn_by_type"]=abn["异常类型"].value_counts()
    s["abn_by_risk"]=abn["风险等级"].value_counts().reindex(["高","中","低"]).fillna(0).astype(int)
    s["abn_by_squad"]=abn["中队"].value_counts().reindex(aq).fillna(0).astype(int)
    s["abn_high"]=abn[abn["风险等级"]=="高"].reset_index(drop=True)
    s["abn_resolved"]=abn["是否化解"].value_counts().to_dict()
    return s

# ── ③ 图表 ──────────────────────────────────────────────
def make_charts(s):
    p={}; C=["#2563eb","#f97316","#10b981","#8b5cf6"]
    def bar(fs,idx,vals,title,yl,key):
        fig,ax=plt.subplots(figsize=fs); clrs=C[:len(idx)]
        bars=ax.bar(idx,vals,color=clrs); ax.set_ylabel(yl); ax.set_title(title)
        top=max(vals) if len(vals) else 1
        for b in bars: ax.text(b.get_x()+b.get_width()/2,b.get_height()+top*.015,f"{b.get_height():.0f}",ha="center",fontsize=10)
        fig.tight_layout(); p[key]=IMG_DIR/f"{key}.png"; fig.savefig(p[key],dpi=130); plt.close(fig)
    sq=s["by_squad"]; bar((6,3.2),sq.index,sq["达标率"],"各中队训练达标率对比","达标率(%)","squad")
    cat=s["by_cat"]; bar((6,3.2),cat.index,cat["平均成绩"],"各训练类别平均考核成绩","平均成绩","cat")
    t=s["abn_by_type"]; bar((6,3.2),t.index,t.values,"异常行为类型分布","起数","abn_type")
    fig,ax=plt.subplots(figsize=(4.2,3.6)); r=s["abn_by_risk"]
    ax.pie(r.values,labels=[f"{k}风险" for k in r.index],autopct="%1.0f%%",colors=["#ef4444","#f97316","#22c55e"],startangle=90)
    ax.set_title("异常行为风险等级分布"); fig.tight_layout(); p["abn_risk"]=IMG_DIR/"abn_risk.png"; fig.savefig(p["abn_risk"],dpi=130); plt.close(fig)
    return p


# ── ④ 模板函数（无 LLM 时降级兜底）────────────────────────
def _t_overview(s):
    return(f"本报告为{META['unit']}{META['period']}{META['cycle']}训练任务与异常行为数据分析报告。本期纳入统计罪犯{s['inmate_count']}名，训练任务记录{s['train_records']}条，异常行为记录{s['abn_total']}起。总体看，平均出勤率{s['avg_att']}%，整体达标率{s['pass_rate']}%，平均考核成绩{s['avg_score']}分；异常行为以{s['abn_by_type'].index[0]}为主，需重点关注高风险个案跟进化解。")
def _t_train_overall(s):
    lv="良好" if s["pass_rate"]>=80 else("一般" if s["pass_rate"]>=60 else"偏低")
    return f"本期参训罪犯共{s['inmate_count']}名，累计记录{s['train_records']}条。平均出勤率{s['avg_att']}%，整体达标率{s['pass_rate']}%，平均考核成绩{s['avg_score']}分，总体完成情况{lv}。"
def _t_train_cat(s):
    c=s["by_cat"]; b=c["达标率"].idxmax(); w=c["达标率"].idxmin()
    return f"分类别看，{b}达标率最高（{c.loc[b,'达标率']}%，均分{c.loc[b,'平均成绩']}分），完成最好；{w}达标率最低（{c.loc[w,'达标率']}%，均分{c.loc[w,'平均成绩']}分），需重点加强。"
def _t_train_squad(s):
    q=s["by_squad"]; b=q["达标率"].idxmax(); w=q["达标率"].idxmin()
    return f"分中队看，{b}达标率最高（{q.loc[b,'达标率']}%）；{w}相对靠后（{q.loc[w,'达标率']}%），建议加强督导。"
def _t_train_trend(s):
    t=s["trend"]; up,fl,dn=t.get("进步",0),t.get("持平",0),t.get("退步",0)
    return f"环比上月，进步{up}项、持平{fl}项、退步{dn}项，整体态势{'稳中向好' if up>=dn else '存在一定波动'}。退步项目应及时分析原因并跟进帮扶。"
def _t_train_focus(s):
    f=s["focus"]
    if not len(f): return "本期未发现训练严重落后的重点人员，整体较为均衡。"
    names="、".join(f"{r['姓名']}（{r['编号']}）" for _,r in f.iterrows())
    return f"经筛查，本期训练重点关注人员共{len(f)}名：{names}。建议安排针对性补训与教育谈话。"
def _t_abn_overall(s):
    t=s["abn_by_type"]; dist="、".join(f"{k}{v}起" for k,v in t.items())
    return f"本期共记录异常行为{s['abn_total']}起，分布：{dist}。{t.index[0]}占比最高，是日常监管主要风险点。"
def _t_abn_risk(s):
    r=s["abn_by_risk"]
    return f"按风险等级：高风险{int(r['高'])}起、中风险{int(r['中'])}起、低风险{int(r['低'])}起。高风险事件须落实「一人一策」，确保管控到位。"
def _t_abn_squad(s):
    q=s["abn_by_squad"]; tp=q.idxmax()
    return f"分中队看，异常行为主要集中在{tp}（{int(q[tp])}起），建议结合训练表现综合研判，加强巡查与心理疏导。"
def _t_abn_high(s):
    h=s["abn_high"]
    if not len(h): return "本期未发生高风险异常行为。"
    _st={"是":"已化解","否":"尚未化解","跟进中":"跟进中"}
    items="；".join(f"{r['姓名']}（{r['编号']}，{r['中队']}）{r['异常类型']}，现{_st.get(r['是否化解'],r['是否化解'])}" for _,r in h.iterrows())
    return f"高风险个案共{len(h)}起：{items}。"
def _t_abn_disposal(s):
    d=s["abn_resolved"]; dn,og,no=d.get("是",0),d.get("跟进中",0),d.get("否",0)
    rt=round(dn/s["abn_total"]*100,1) if s["abn_total"] else 0
    return f"处置化解：已化解{dn}起、跟进中{og}起、未化解{no}起，化解率{rt}%。跟进中及未化解事件应明确责任人、限时销号。"
def _t_conclusion(s):
    wt=s["by_squad"]["达标率"].idxmin(); wa=s["abn_by_squad"].idxmax()
    lk="，且二者集中在同一中队，需作为本期重点管理对象" if wt==wa else "，训练落后与异常高发分布在不同中队，应分类施策"
    return f"综合研判：本期{META['unit']}训练达标率{s['pass_rate']}%，异常{s['abn_total']}起，化解工作总体有序。训练薄弱中队为{wt}，异常集中中队为{wa}{lk}。总体可控，局部需加强。"
def _t_suggestion(s):
    wc=s["by_cat"]["达标率"].idxmin(); hn=len(s["abn_high"])
    return f"建议：一是针对{wc}等薄弱科目制定专项提升计划；二是落实{hn}起高风险个案「一人一策」管控；三是对退步及不达标人员开展补训与教育谈话；四是健全数据采集机制，为后续分析提供支撑。"


# ── ⑤ 生成函数（LLM 模式，各章节打包结构化数据）────────────
def n_overview(s):
    data=(f"单位：{META['unit']}  期间：{META['period']}  周期：{META['cycle']}\n"
          f"纳入统计罪犯：{s['inmate_count']}名，训练任务记录：{s['train_records']}条，异常行为记录：{s['abn_total']}起\n"
          f"平均出勤率：{s['avg_att']}%，整体达标率：{s['pass_rate']}%（达标标准：成绩≥{s['pass_score']}分且出勤率≥{s['pass_att']}%）\n"
          f"平均考核成绩：{s['avg_score']}分\n"
          f"异常行为主要类型：{s['abn_by_type'].index[0]}（{s['abn_by_type'].iloc[0]}起，占比最高）")
    return _gen("overview",data,_t_overview,s)

def n_train_overall(s):
    lv="良好" if s["pass_rate"]>=80 else("一般" if s["pass_rate"]>=60 else"偏低")
    data=(f"参训罪犯：{s['inmate_count']}名，训练记录：{s['train_records']}条\n"
          f"平均出勤率：{s['avg_att']}%\n"
          f"达标率：{s['pass_rate']}%（达标标准：成绩≥{s['pass_score']}分且出勤率≥{s['pass_att']}%）\n"
          f"平均考核成绩：{s['avg_score']}分\n"
          f"整体完成水平评价：{lv}")
    return _gen("train_overall",data,_t_train_overall,s)

def n_train_cat(s):
    c=s["by_cat"]; b=c["达标率"].idxmax(); w=c["达标率"].idxmin()
    rows="\n".join(f"  {idx}：达标率{row['达标率']}%，平均出勤率{row['平均出勤率']}%，平均成绩{row['平均成绩']}分" for idx,row in c.iterrows())
    data=(f"各训练类别完成情况：\n{rows}\n"
          f"达标率最高类别：{b}（{c.loc[b,'达标率']}%，均分{c.loc[b,'平均成绩']}分）\n"
          f"达标率最低类别：{w}（{c.loc[w,'达标率']}%，均分{c.loc[w,'平均成绩']}分）")
    return _gen("train_cat",data,_t_train_cat,s)

def n_train_squad(s):
    q=s["by_squad"]; b=q["达标率"].idxmax(); w=q["达标率"].idxmin()
    rows="\n".join(f"  {idx}：达标率{row['达标率']}%，平均出勤率{row['平均出勤率']}%，平均成绩{row['平均成绩']}分" for idx,row in q.iterrows())
    data=(f"各中队训练完成情况：\n{rows}\n"
          f"达标率最高中队：{b}（{q.loc[b,'达标率']}%）\n"
          f"达标率最低中队：{w}（{q.loc[w,'达标率']}%），建议重点督导")
    return _gen("train_squad",data,_t_train_squad,s)

def n_train_trend(s):
    t=s["trend"]; up,fl,dn=t.get("进步",0),t.get("持平",0),t.get("退步",0)
    data=(f"环比上月变化统计：\n"
          f"  进步：{up}项\n  持平：{fl}项\n  退步：{dn}项\n"
          f"进退比：{up}/{dn}（进步/退步）\n"
          f"整体趋势：{'稳中向好' if up>=dn else '存在一定波动，需关注退步项目原因'}")
    return _gen("train_trend",data,_t_train_trend,s)

def n_train_focus(s):
    f=s["focus"]; ft=s.get("focus_threshold",70)
    if not len(f): return "本期未发现训练严重落后的重点人员，整体较为均衡。"
    rows="\n".join(f"  {r['姓名']}（{r['编号']}，{r['中队']}）：平均成绩{r['平均成绩']}分，出勤率{r['平均出勤率']}%，不达标{r['不达标项']}项，退步{r['退步项']}项" for _,r in f.iterrows())
    data=(f"筛选条件：平均成绩<{ft}分或不达标项≥2项\n"
          f"重点关注人员共{len(f)}名：\n{rows}")
    return _gen("train_focus",data,_t_train_focus,s)

def n_abn_overall(s):
    t=s["abn_by_type"]
    dist="\n".join(f"  {k}：{v}起" for k,v in t.items())
    data=(f"本期异常行为总计：{s['abn_total']}起\n"
          f"各类型分布：\n{dist}\n"
          f"占比最高类型：{t.index[0]}（{t.iloc[0]}起）")
    return _gen("abn_overall",data,_t_abn_overall,s)

def n_abn_risk(s):
    r=s["abn_by_risk"]
    high_pct=round(int(r['高'])/s['abn_total']*100,1) if s['abn_total'] else 0
    data=(f"异常行为风险等级分布：\n"
          f"  高风险：{int(r['高'])}起\n"
          f"  中风险：{int(r['中'])}起\n"
          f"  低风险：{int(r['低'])}起\n"
          f"高风险占比：{high_pct}%")
    return _gen("abn_risk",data,_t_abn_risk,s)

def n_abn_squad(s):
    q=s["abn_by_squad"]; tp=q.idxmax()
    rows="\n".join(f"  {k}：{v}起" for k,v in q.items())
    data=(f"各中队异常行为起数：\n{rows}\n"
          f"异常最集中中队：{tp}（{int(q[tp])}起）")
    return _gen("abn_squad",data,_t_abn_squad,s)

def n_abn_high(s):
    h=s["abn_high"]
    if not len(h): return "本期未发生高风险异常行为。"
    _st={"是":"已化解","否":"尚未化解","跟进中":"跟进中"}
    rows="\n".join(f"  {r['姓名']}（{r['编号']}，{r['中队']}）：{r['异常类型']}，发生于{r['发生日期']}，当前状态：{_st.get(r['是否化解'],r['是否化解'])}" for _,r in h.iterrows())
    data=f"高风险异常个案共{len(h)}起：\n{rows}"
    return _gen("abn_high",data,_t_abn_high,s)

def n_abn_disposal(s):
    d=s["abn_resolved"]; dn,og,no=d.get("是",0),d.get("跟进中",0),d.get("否",0)
    rt=round(dn/s["abn_total"]*100,1) if s["abn_total"] else 0
    data=(f"处置化解情况：\n"
          f"  已化解：{dn}起\n  跟进中：{og}起\n  未化解：{no}起\n"
          f"化解率：{rt}%\n"
          f"仍需跟进（跟进中+未化解）：{og+no}起")
    return _gen("abn_disposal",data,_t_abn_disposal,s)

def n_conclusion(s):
    wt=s["by_squad"]["达标率"].idxmin(); wa=s["abn_by_squad"].idxmax()
    d=s["abn_resolved"]; rt=round(d.get("是",0)/s["abn_total"]*100,1) if s["abn_total"] else 0
    data=(f"综合情况：\n"
          f"  单位：{META['unit']}  期间：{META['period']}\n"
          f"  整体训练达标率：{s['pass_rate']}%  平均成绩：{s['avg_score']}分\n"
          f"  本期异常行为：{s['abn_total']}起，化解率：{rt}%\n"
          f"  训练达标率最低中队：{wt}（达标率{s['by_squad'].loc[wt,'达标率']}%）\n"
          f"  异常行为最多中队：{wa}（{int(s['abn_by_squad'][wa])}起）\n"
          f"  两者是否为同一中队：{'是，需列为本期重点管理对象' if wt==wa else '否，应分类施策分别管控'}")
    return _gen("conclusion",data,_t_conclusion,s)

def n_suggestion(s):
    wc=s["by_cat"]["达标率"].idxmin(); hn=len(s["abn_high"])
    t=s["trend"]; dn_cnt=t.get("退步",0); fn=len(s["focus"])
    rt=round(s["abn_resolved"].get("是",0)/s["abn_total"]*100,1) if s["abn_total"] else 0
    data=(f"需改进情况汇总：\n"
          f"  训练达标率最低类别：{wc}（达标率{s['by_cat'].loc[wc,'达标率']}%）\n"
          f"  高风险异常个案：{hn}起\n"
          f"  训练退步项目：{dn_cnt}项\n"
          f"  重点关注人员：{fn}名\n"
          f"  异常化解率：{rt}%（仍有{s['abn_total']-s['abn_resolved'].get('是',0)}起待处置）")
    return _gen("suggestion",data,_t_suggestion,s)


# ── ⑥ Word 输出 ─────────────────────────────────────────
_ALL_SEC=["train_overall","train_cat","train_squad","train_trend","train_focus",
          "abn_overall","abn_risk","abn_squad","abn_high","abn_disposal","conclusion","suggestion"]

def _s(run,name,size,bold=False,color=None):
    run.font.name=name; run.font.size=Pt(size); run.font.bold=bold
    if color: run.font.color.rgb=color
    run._element.rPr.rFonts.set(qn("w:eastAsia"),name)
def add_title(doc,t):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; _s(p.add_run(t),F_TITLE,22,True)
def add_h1(doc,t):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); _s(p.add_run(t),F_HEAD,15,True)
def add_h2(doc,t):
    p=doc.add_paragraph(); _s(p.add_run(t),F_HEAD,13,True)
def add_body(doc,text):
    # 直接写入：生成模式下 LLM 已完成内容创作，无需二次加工
    p=doc.add_paragraph()
    p.paragraph_format.first_line_indent=Pt(28); p.paragraph_format.line_spacing=1.5
    _s(p.add_run(text),F_BODY,14)
def add_img(doc,path,width=5.8):
    doc.add_picture(str(path),width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
def add_df_table(doc,df,idx_label=None):
    cols=([idx_label] if idx_label else[])+list(df.columns)
    t=doc.add_table(rows=1,cols=len(cols)); t.style="Light Grid Accent 1"
    for i,c in enumerate(cols): _s(t.rows[0].cells[i].paragraphs[0].add_run(str(c)),F_HEAD,10.5,True)
    for ix,row in df.iterrows():
        cells=t.add_row().cells; vals=([ix] if idx_label else[])+list(row.values)
        for i,v in enumerate(vals): _s(cells[i].paragraphs[0].add_run(f"{v:.1f}" if isinstance(v,float) else str(v)),F_BODY,10.5)

def build_report(stats,charts,sections=None,classification="内部资料"):
    s=stats; SEC={k:True for k in _ALL_SEC}
    if sections: SEC.update(sections)
    doc=Document()
    cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    _s(cp.add_run(f"【{classification}】"),F_BODY,11,color=RGBColor(0xC0,0,0) if classification in("保密","机密") else RGBColor(0x60,0x60,0x60))
    add_title(doc,f"{META['unit']}{META['period']}罪犯训练任务与异常行为数据分析报告")
    sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    pts=[f"报送单位：{META['unit']}",f"呈报：{META['parent']}"]
    if META.get("author"): pts.append(f"编制：{META['author']}")
    if META.get("reviewer"): pts.append(f"审核：{META['reviewer']}")
    pts.append(f"生成时间：{datetime.now():%Y-%m-%d}")
    _s(sub.add_run("    ".join(pts)),F_BODY,10.5)
    add_h1(doc,"一、报告概述"); add_body(doc,n_overview(s))
    add_h1(doc,"二、训练任务完成情况")
    if SEC["train_overall"]: add_h2(doc,"（一）总体完成情况"); add_body(doc,n_train_overall(s))
    if SEC["train_cat"]: add_h2(doc,"（二）分类别分析"); add_body(doc,n_train_cat(s)); add_df_table(doc,s["by_cat"],"训练类别"); add_img(doc,charts["cat"])
    if SEC["train_squad"]: add_h2(doc,"（三）分中队对比"); add_body(doc,n_train_squad(s)); add_df_table(doc,s["by_squad"],"中队"); add_img(doc,charts["squad"])
    if SEC["train_trend"]: add_h2(doc,"（四）环比趋势"); add_body(doc,n_train_trend(s))
    if SEC["train_focus"]: add_h2(doc,"（五）重点关注人员"); add_body(doc,n_train_focus(s))
    if SEC["train_focus"] and len(s["focus"]): add_df_table(doc,s["focus"])
    add_h1(doc,"三、异常行为情况")
    if SEC["abn_overall"]: add_h2(doc,"（一）总体态势"); add_body(doc,n_abn_overall(s)); add_img(doc,charts["abn_type"])
    if SEC["abn_risk"]: add_h2(doc,"（二）风险等级分布"); add_body(doc,n_abn_risk(s)); add_img(doc,charts["abn_risk"],width=3.6)
    if SEC["abn_squad"]: add_h2(doc,"（三）分中队分布"); add_body(doc,n_abn_squad(s))
    if SEC["abn_high"]: add_h2(doc,"（四）高风险个案"); add_body(doc,n_abn_high(s))
    if SEC["abn_disposal"]: add_h2(doc,"（五）处置与化解情况"); add_body(doc,n_abn_disposal(s))
    if SEC["conclusion"]: add_h1(doc,"四、综合研判与结论"); add_body(doc,n_conclusion(s))
    if SEC["suggestion"]: add_h1(doc,"五、下一步工作建议"); add_body(doc,n_suggestion(s))
    out=OUT_DIR/f"{META['unit']}{META['period']}训练与异常行为数据分析报告.docx"
    doc.save(out); return out

def main():
    print("大脑模式：",f"Mimo/{MIMO_MODEL}（大模型生成）" if USE_LLM else "模板模式（无Key）")
    train,abn=load_data(); stats=analyze(train,abn); charts=make_charts(stats); out=build_report(stats,charts)
    print("完成 ->",out)
if __name__=="__main__": main()
